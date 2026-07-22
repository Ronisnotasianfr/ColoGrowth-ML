"""
build_paper.py - Assemble the Word and LaTeX research paper from leakage-fixed pipeline outputs.
Updated to include all Phase 3 polish improvements (benchmarking, subgroup interactions, split justifications).

Usage:
    python paper/build_paper.py --dataset geo
"""

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from paper_metrics import (
    project_root,
    load_leakage_fixed_metrics,
    load_dataset_stats,
    build_abstract,
    build_methods_leakage_paragraph,
    build_results_opening,
    build_results_closing,
    build_discussion_paragraph,
    metrics_table_rows,
    dataset_table_rows,
    build_hyperparameters_table,
    build_top_genes_table,
    build_nnt_table,
    build_subgroups_table,
    build_cox_table,
    build_sensitivity_table,
    build_benchmarking_table,
    build_methods_split_justification,
    build_discussion_pathway_expansion,
    build_introduction_p1,
    build_introduction_p2,
    build_methods_p1,
    build_methods_p2,
    build_interpretation_p1,
    build_interpretation_p2,
    build_interpretation_p3,
    build_clinical_validation_p1,
    build_clinical_validation_p2,
    build_clinical_validation_p3,
    build_cox_paragraph,
    build_sensitivity_p1,
    build_discussion_benchmarking_intro,
)

TITLE = "Cross-Platform Colon Cancer Proliferation Classification via Leakage-Free ML"
SUBTITLE = "Independent computational biology research report prepared for peer-reviewed journal submission"
AUTHOR = "Rohan Saindane"
DATE = "June 2026"

INK = RGBColor(20, 31, 43)
BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(90, 90, 90)
LIGHT_FILL = "F4F6F9"
BORDER = "D9DEE8"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, size=9.0, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def style_run(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", after=7, before=0, align=None, size=10.6, bold=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        style_run(r, size=size, color=INK, bold=bold)
    return p


def add_ascii_schematic(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8.0)
    run.font.color.rgb = INK
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    style_run(r, size=8.8, italic=True, color=GRAY)


def add_figure(doc, path, caption, width=5.45):
    if not Path(path).exists():
        add_para(doc, f"[Figure unavailable: {path}]", size=9.0, color=GRAY)
        add_caption(doc, caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_page_break(doc):
    doc.add_page_break()


def apply_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = INK

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 12, 5),
        ("Heading 2", 12.5, BLUE, 8, 4),
        ("Heading 3", 11.2, RGBColor(45, 75, 105), 6, 3),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("ColoGrowth-ML | Leakage-free cross-platform validation study")
    style_run(r, size=8.5, color=GRAY)


def add_title_page(doc, metrics, stats):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(TITLE)
    style_run(r, size=18, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(SUBTITLE)
    style_run(r, size=11.5, italic=True, color=GRAY)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = [
        ("Prepared by", AUTHOR),
        ("Date", DATE),
        ("Dataset", stats['dataset'].upper()),
        ("Status", "Leakage-corrected journal submission draft"),
    ]
    for i, (k, v) in enumerate(labels):
        set_cell_text(meta.cell(i, 0), k, bold=True, size=8.8, color=GRAY)
        set_cell_text(meta.cell(i, 1), v, size=8.8, color=INK)
    set_table_borders(meta, "FFFFFF")

    add_heading(doc, "Abstract", 1)
    add_para(doc, build_abstract(metrics, stats), size=10.4)

    add_heading(doc, "1. Introduction", 1)
    add_para(doc, build_introduction_p1(), size=10.4)
    add_para(doc, build_introduction_p2(), size=10.4)


def add_methods(doc, stats):
    add_heading(doc, "2. Materials and Methods", 1)
    add_para(doc, build_methods_p1(stats), size=10.4)

    data_table = doc.add_table(rows=1, cols=4)
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = data_table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ["Dataset file", "Samples", "Features/columns", "Class balance"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.8, color=INK)
    for row in dataset_table_rows(stats):
        cells = data_table.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.6)
    set_table_borders(data_table)
    add_caption(doc, "Table 1. Processed data files used for this leakage-corrected report.")

    # Task 3.3 - Three-way validation split justification
    add_para(doc, build_methods_split_justification(), size=10.4)
    add_para(doc, "The three-way validation data-flow schematic is detailed in Figure 7.", size=10.4)
    add_ascii_schematic(
        doc,
        "                     [ GEO Cohort (n=585) ]\n"
        "                               |\n"
        "                 +-------------+-------------+\n"
        "                 v                           v\n"
        "        GEO-Train (80%, n=468)      GEO-Holdout (20%, n=117)\n"
        "                 |                           |\n"
        "                 +-> [Feature Selection]     +-> [Evaluate Model]\n"
        "                 +-> [GridSearchCV Tuning]   |\n"
        "                 +-> [Fit Model Coefficients]|\n"
        "                                             |\n"
        "          +----------------------------------+\n"
        "          v                                   v\n"
        "  [TCGA Cohort (n=322)]            [CPTAC Cohort (n=105)]\n"
        "          |                                   |\n"
        "    +-----+-----+                       +----+------+\n"
        "    v           v                       v           v\n"
        "  Calib (50%)  Eval (50%)           Calib (50%)  Eval (50%)"
    )
    add_caption(doc, "Figure 7. Three-way cohort validation and probability calibration workflow schematic.")

    add_para(doc, build_methods_p2(), size=10.4)
    
    # Target Leakage prevention
    add_para(doc, build_methods_leakage_paragraph(), size=10.4)

    # Table S2: Hyperparameters
    add_para(doc, "GridSearchCV parameter search spaces and optimal configurations are shown in Table 2.", size=10.4)
    hp_table = doc.add_table(rows=1, cols=5)
    hp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = hp_table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ["Model", "Hyperparameter", "Search Range", "Selected", "Notes"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.2, color=INK)
    for row in build_hyperparameters_table():
        cells = hp_table.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.0)
    set_table_borders(hp_table)
    add_caption(doc, "Table 2. Hyperparameter optimization search ranges and selected settings.")


def add_results(doc, metrics, results_dir):
    add_heading(doc, "3. Results", 1)
    add_para(doc, build_results_opening(metrics), size=10.4)

    # Table 3: Model Metrics
    perf = doc.add_table(rows=1, cols=4)
    perf.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Model", "CV ROC-AUC (mean +/- std)", "Holdout Accuracy (95% CI)", "Holdout ROC-AUC (95% CI)"]
    for cell, text in zip(perf.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.2, color=INK)
    set_repeat_table_header(perf.rows[0])
    for row in metrics_table_rows(metrics):
        cells = perf.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.0)
    set_table_borders(perf)
    add_caption(doc, "Table 3. Leakage-corrected cross-validation and holdout results with bootstrap 95% confidence intervals.")

    # Calibration comparison curve
    add_figure(
        doc,
        results_dir / "calibration_comparison_curves.png",
        "Figure 1. Calibration curves showing observed positive fraction vs. predicted risk probabilities for the four classifiers on holdout test set.",
        width=5.2,
    )
    
    # External validation and calibration text
    add_para(doc, build_results_closing(metrics), size=10.4)


def add_interpretation(doc, results_dir):
    add_heading(doc, "4. Interpretation and Biological Readout", 1)
    add_para(doc, build_interpretation_p1(), size=10.4)

    # Table 4: Feature Genes
    genes_tab = doc.add_table(rows=1, cols=4)
    genes_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(genes_tab.rows[0].cells, ["Rank", "Gene Symbol", "ANOVA F-Score", "CV Selection Frequency"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.5, color=INK)
    set_repeat_table_header(genes_tab.rows[0])
    for row in build_top_genes_table():
        cells = genes_tab.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.2)
    set_table_borders(genes_tab)
    add_caption(doc, "Table 4. Top selected transcriptomic feature genes ranked by ANOVA F-score.")

    add_figure(
        doc,
        results_dir / "shap_summary_random_forest.png",
        "Figure 2. Random forest SHAP summary showing feature impact (red/high, blue/low expression) on proliferation class prediction.",
        width=5.35,
    )
    
    # Pathway Enrichment
    add_para(doc, build_interpretation_p2(), size=10.4)
    add_figure(
        doc,
        results_dir / "pathway_enrichment.png",
        "Figure 3. Enriched biological pathways (FDR < 0.05) representing transcriptomic cascade signatures downstream of cancer cell division.",
        width=5.35,
    )
    
    # Clinical DCA
    add_para(doc, build_interpretation_p3(), size=10.4)
    add_figure(
        doc,
        results_dir / "clinical_dca.png",
        "Figure 4. Clinical Decision Curve Analysis comparing Net Benefit of model-guided stratification vs. default intervention strategies.",
        width=5.2,
    )
    
    # Table S3: NNT
    nnt_tab = doc.add_table(rows=1, cols=7)
    nnt_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(nnt_tab.rows[0].cells, ["Model", "Threshold", "Sensitivity", "Specificity", "PPV", "NPV", "NNT"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.2, color=INK)
    set_repeat_table_header(nnt_tab.rows[0])
    for row in build_nnt_table():
        cells = nnt_tab.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.0)
    set_table_borders(nnt_tab)
    add_caption(doc, "Table 5. Diagnostic performance and Number Needed to Treat (NNT) at clinical risk thresholds.")


def add_clinical_validation(doc, results_dir):
    add_heading(doc, "5. Demographic Subgroups & Prognostic Validation", 1)
    
    # Subgroups text
    add_para(doc, build_clinical_validation_p1(), size=10.4)
    
    # Table S4: Subgroups (7 columns now!)
    sub_tab = doc.add_table(rows=1, cols=7)
    sub_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Subgroup", "N", "Accuracy", "ROC-AUC", "ROC-AUC 95% CI", "Interaction p-value", "Interaction 95% CI"]
    for cell, text in zip(sub_tab.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.2, color=INK)
    set_repeat_table_header(sub_tab.rows[0])
    for row in build_subgroups_table():
        cells = sub_tab.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.0)
    set_table_borders(sub_tab)
    add_caption(doc, "Table 6. Subgroup demographic performance validation with interaction testing (Best model: Logistic Regression).")

    # Kaplan-Meier
    add_para(doc, build_clinical_validation_p2(), size=10.4)
    
    add_figure(
        doc,
        results_dir / "kaplan_meier_geo.png",
        "Figure 5. Kaplan-Meier overall survival curves comparing predicted high vs. low proliferation cohorts in GEO.",
        width=5.1,
    )
    
    add_figure(
        doc,
        results_dir / "kaplan_meier_stage_stratified.png",
        "Figure 6. Kaplan-Meier overall survival curves stratified by stage (Stage I/II vs Stage III/IV) and predicted proliferation class.",
        width=5.4,
    )
    
    # Cox PH Table
    add_para(doc, build_clinical_validation_p3(), size=10.4)
    
    cox_tab = doc.add_table(rows=1, cols=5)
    cox_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(cox_tab.rows[0].cells, ["Predictor / Covariate", "Coefficient", "Hazard Ratio (HR)", "95% CI", "p-value"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.5, color=INK)
    set_repeat_table_header(cox_tab.rows[0])
    for row in build_cox_table():
        cells = cox_tab.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.2)
    set_table_borders(cox_tab)
    add_caption(doc, "Table 7. Multivariate Cox Proportional Hazards survival analysis summary.")


def add_sensitivity_section(doc):
    add_heading(doc, "6. Pre-Processing Sensitivity & Robustness Analysis", 1)
    
    add_para(doc, build_sensitivity_p1(), size=10.4)
    
    # Table S6: Sensitivity
    sens_tab = doc.add_table(rows=1, cols=5)
    sens_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["SelectKBest k", "Holdout ROC-AUC (k)", "Variance Threshold (VT)", "Features Passed VT", "Holdout ROC-AUC (VT)"]
    for cell, text in zip(sens_tab.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.2, color=INK)
    set_repeat_table_header(sens_tab.rows[0])
    for row in build_sensitivity_table():
        cells = sens_tab.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.0)
    set_table_borders(sens_tab)
    add_caption(doc, "Table 8. Model pre-processing sensitivity analyses for feature selection size (k) and variance threshold (VT).")


def add_discussion(doc):
    add_heading(doc, "7. Discussion", 1)
    add_para(doc, build_discussion_paragraph(), size=10.4)

    # Task 3.1 - Benchmarking Table
    add_para(doc, build_discussion_benchmarking_intro(), size=10.4)
    bench_tab = doc.add_table(rows=1, cols=7)
    bench_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Study", "Year", "Cohort/Platform", "N", "AUC/Accuracy", "Leakage-controlled?", "Cross-platform validated?"]
    for cell, text in zip(bench_tab.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=8.2, color=INK)
    set_repeat_table_header(bench_tab.rows[0])
    for row in build_benchmarking_table():
        cells = bench_tab.add_row().cells
        for c, text in zip(cells, row):
            set_cell_text(c, text, size=8.0)
    set_table_borders(bench_tab)
    add_caption(doc, "Table 9. Performance and methodological comparisons with published signatures.")

    # Task 3.4 - Biological Mechanism Discussion Expansion
    add_para(doc, build_discussion_pathway_expansion(), size=10.4)

    add_heading(doc, "8. Limitations and Future Directions", 1)
    
    add_para(doc, "LIMITATIONS:", size=10.4, bold=True)
    limitations = [
        "Sample size: GEO GSE39582 (n=585) is moderate. CPTAC-COAD (n=105) has only 7 survival events, limiting power.",
        "Target binarization at the median is standard but arbitrary. Continuous risk scores might work better.",
        "Microarray and RNA-seq have different dynamic ranges, requiring post-hoc calibration to restore accuracy.",
        "SHAP scores reflect correlation, not causation."
    ]
    for lim in limitations:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(lim)
        style_run(r, size=10.2, color=INK)
        
    add_para(doc, "FUTURE WORK:", size=10.4, bold=True)
    future = [
        "Prospective validation of the Top-3 Ensemble on new COAD biopsy cohorts.",
        "qPCR knockdown of top SHAP genes (RPS3, RPS11) to confirm their role in growth regulation.",
        "Integration of CNVs and somatic mutation data into the feature space."
    ]
    for fut in future:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(fut)
        style_run(r, size=10.2, color=INK)


def add_references(doc):
    add_heading(doc, "References", 1)
    refs = [
        "Marisa, L. et al. Gene expression classification of colon cancer into molecular subtypes: characterization, validation, and prognostic value. PLoS Medicine, 2013. DOI: 10.1371/journal.pmed.1001453",
        "Whitfield, M. L. et al. Identification of genes periodically expressed in the human cell cycle and their expression in tumors. Molecular Biology of the Cell, 2002. DOI: 10.1091/mbc.02-02-0030",
        "Lundberg, S. M. and Lee, S.-I. A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 2017.",
        "The Cancer Genome Atlas Research Network. TCGA Colon Adenocarcinoma data resource, accessed through the NCI Genomic Data Commons.",
        "National Center for Biotechnology Information Gene Expression Omnibus. GSE39582 dataset record.",
        "Zeng, D.-T. et al. Prognostic role of Ki-67 in colorectal carcinoma: Development and evaluation of machine learning prediction models. World Journal of Clinical Oncology, 2025. DOI: 10.5306/wjco.v16.i8.107306",
        "Agesen, T. H. et al. ColoGuideEx: a robust gene classifier specific for stage II colorectal cancer prognosis. Gut, 2012. DOI: 10.1136/gutjnl-2011-301179",
        "O'Connell, M. J. et al. Relationship between tumor gene expression and recurrence in four independent studies of patients with stage II/III colon cancer treated with surgery alone or surgery plus adjuvant fluorouracil plus leucovorin. Journal of Clinical Oncology, 2010. DOI: 10.1200/JCO.2010.28.9538",
        "Langston, L. D. et al. Mcm10 promotes rapid isomerization of CMG-DNA for replisome bypass of lagging strand DNA blocks. eLife, 2017. DOI: 10.7554/eLife.29118",
        "Bharadwaj, R., Qi, W., and Yu, H. Identification of two novel components of the human NDC80 kinetochore complex. Journal of Biological Chemistry, 2004. DOI: 10.1074/jbc.M310224200",
        "Seipold, S. et al. Non-SMC condensin I complex proteins control chromosome segregation and survival of proliferating cells in the zebrafish neural retina. BMC Developmental Biology, 2009. DOI: 10.1186/1471-213X-9-40",
        "Overmeer, R. M. et al. Replication factor C recruits DNA polymerase delta to sites of nucleotide excision repair but is not required for PCNA recruitment. Molecular and Cellular Biology, 2010. DOI: 10.1128/MCB.00285-10"
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ref)
        style_run(r, size=9.8, color=INK)

    add_heading(doc, "Data and Code Availability", 1)
    add_para(
        doc,
        "GEO GSE39582 is available at: https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE39582. "
        "TCGA-COAD is available at UCSC Xena: https://xenabrowser.net/. "
        "CPTAC-COAD data is available at: https://cptac-data-portal.georgetown.edu/. "
        "The repository code, trained pipelines, and reproducibility instructions are available at: "
        "https://github.com/Ronisnotasianfr/ColoGrowth-ML.",
        size=10.2,
    )

    add_heading(doc, "Ethical Considerations and AI Disclosure", 1)
    add_para(
        doc,
        "Secondary analysis of de-identified public datasets did not require institutional review board (IRB) "
        "approval. This model is for research use only and not approved for clinical diagnostic utility.",
        size=10.2,
    )
    add_heading(doc, "AI Assistance", 2)
    add_para(
        doc,
        "Claude (Anthropic) was used as a coding assistant during implementation. All scientific decisions, "
        "study design, data interpretation, and conclusions are the author's own. Full prompt logs are "
        "available in the project repository.",
    )


def build_latex(metrics, stats, out_path):
    abstract_text = build_abstract(metrics, stats)
    methods_leakage = build_methods_leakage_paragraph()
    methods_split = build_methods_split_justification()
    results_opening = build_results_opening(metrics)
    results_closing = build_results_closing(metrics)
    discussion = build_discussion_paragraph()
    discussion_pathway = build_discussion_pathway_expansion()

    t1_rows = []
    for row in dataset_table_rows(stats):
        t1_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} \\\\")
    t1_content = "\n".join(t1_rows)

    t2_rows = []
    for row in metrics_table_rows(metrics):
        t2_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} \\\\")
    t2_content = "\n".join(t2_rows)

    # Advanced tables
    hp_rows = []
    for row in build_hyperparameters_table():
        hp_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} & {row[4]} \\\\")
    hp_content = "\n".join(hp_rows)

    genes_rows = []
    for row in build_top_genes_table():
        genes_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} \\\\")
    genes_content = "\n".join(genes_rows)

    nnt_rows = []
    for row in build_nnt_table():
        nnt_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} & {row[4]} & {row[5]} & {row[6]} \\\\")
    nnt_content = "\n".join(nnt_rows)

    sub_rows = []
    for row in build_subgroups_table():
        sub_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} & {row[4]} & {row[5]} & {row[6]} \\\\")
    sub_content = "\n".join(sub_rows)

    cox_rows = []
    for row in build_cox_table():
        cox_rows.append(f"  {row[0].replace('_', '\\_')} & {row[1]} & {row[2]} & {row[3]} & {row[4]} \\\\")
    cox_content = "\n".join(cox_rows)

    cox_paragraph = build_cox_paragraph()

    sens_rows = []
    for row in build_sensitivity_table():
        sens_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} & {row[4]} \\\\")
    sens_content = "\n".join(sens_rows)

    bench_rows = []
    for row in build_benchmarking_table():
        bench_rows.append(f"  {row[0]} & {row[1]} & {row[2]} & {row[3]} & {row[4]} & {row[5]} & {row[6]} \\\\")
    bench_content = "\n".join(bench_rows)

    tex_content = f"""\\documentclass[11pt]{{article}}
\\usepackage[margin=0.8in]{{geometry}}
\\usepackage{{graphicx}}
\\usepackage{{booktabs}}
\\usepackage{{array}}
\\usepackage{{times}}
\\usepackage{{caption}}
\\usepackage{{float}}
\\usepackage[hidelinks]{{hyperref}}
\\graphicspath{{{{../}}{{./}}}}

\\title{{\\textbf{{{TITLE}}}}}
\\author{{{AUTHOR}\\\\Independent Research Project}}
\\date{{{DATE}}}

\\begin{{document}}
\\maketitle

\\begin{{abstract}}
{abstract_text}
\\end{{abstract}}

\\section{{Introduction}}
Cell proliferation correlates with survival, recurrence, and chemotherapy response in colon adenocarcinoma. Clinicians estimate proliferation through Ki-67 staining or staging. These methods have inter-observer variability and miss broader transcriptomic changes from cell-cycle deregulation. Machine learning on gene expression data could provide an automated alternative.

We trained classifiers on microarray data to predict high vs. low proliferation from gene expression and clinical covariates. The ten cell-cycle genes that define the target were removed from features before training. We compared Logistic Regression, Random Forest, XGBoost, and a Multilayer Perceptron using nested CV on microarray data with external validation on an independent RNA-seq cohort.

\\section{{Materials and Methods}}
The {stats['dataset'].upper()} processed dataset contains {stats['n_samples']} samples and {stats['n_features']} features after signature-gene removal. Clinical covariates include age, sex, and stage. The binary target is balanced with {stats['class_balance']}. An 80/20 stratified split produced a training pool of {stats['train_n']} samples and a holdout test set of {stats['test_n']} samples.

\\begin{{table}}[H]
\\centering
\\caption{{Processed data files available in the project at the time of this report.}}
\\begin{{tabular}}{{lccc}}
\\toprule
\\textbf{{Dataset file}} & \\textbf{{Samples}} & \\textbf{{Features/columns}} & \\textbf{{Class balance}} \\\\
\\midrule
{t1_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

{methods_split}

\\begin{{figure}}[H]
\\centering
\\begin{{verbatim}}
                     [ GEO Cohort (n=585) ]
                               │
                 ┌─────────────┴─────────────┐
                 ▼                           ▼
        GEO-Train (80%, n=468)      GEO-Holdout (20%, n=117)
                 │                           │
                 ├─► [Feature Selection]     ├─► [Evaluate Model]
                 ├─► [GridSearchCV Tuning]   │
                 └─► [Fit Model Coefficients]│
                                             │
          ┌──────────────────────────────────┘
          ▼                                   ▼
  [TCGA Cohort (n=322)]            [CPTAC Cohort (n=105)]
          │                                   │
    ┌─────┴─────┐                       ┌─────┴──────┐
    ▼           ▼                       ▼            ▼
  Calib (50%)  Eval (50%)           Calib (50%)  Eval (50%)
\\end{{verbatim}}
\\caption{{Three-way cohort validation and probability calibration workflow schematic.}}
\\end{{figure}}

Four classifiers: logistic regression, random forest, XGBoost, and a multilayer perceptron. Each was wrapped in an sklearn Pipeline (standardization, variance filtering, SelectKBest, classifier). Five-fold stratified CV was run on the training pool with fold-local preprocessing, GridSearchCV tuning, and holdout evaluation.

{methods_leakage}

\\begin{{table}}[H]
\\centering
\\caption{{Hyperparameter optimization search ranges and selected settings.}}
\\begin{{tabular}}{{lllll}}
\\toprule
\\textbf{{Model}} & \\textbf{{Hyperparameter}} & \\textbf{{Search Range}} & \\textbf{{Selected}} & \\textbf{{Notes}} \\\\
\\midrule
{hp_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

\\section{{Results}}
{results_opening}

\\begin{{table}}[H]
\\centering
\\caption{{Cross-validation and holdout results from the evaluation run.}}
\\begin{{tabular}}{{lccc}}
\\toprule
\\textbf{{Model}} & \\textbf{{CV ROC-AUC}} & \\textbf{{Holdout Accuracy (95\\% CI)}} & \\textbf{{Holdout ROC-AUC (95\\% CI)}} \\\\
\\midrule
{t2_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width=0.68\\linewidth]{{results/calibration_comparison_curves.png}}
\\caption{{Calibration curves showing observed positive fraction vs. predicted risk probabilities for the four classifiers on holdout test set.}}
\\end{{figure}}

{results_closing}

\\section{{Interpretation and Biological Readout}}
SHAP summaries were generated from pipeline-transformed, leakage-free features. Because the ten signature genes used to define the proliferation label were removed before training, top-ranked features should be interpreted as candidate biological or clinical correlates rather than label-construction artifacts.

\\begin{{table}}[H]
\\centering
\\caption{{Top selected transcriptomic feature genes ranked by ANOVA F-score.}}
\\begin{{tabular}}{{llll}}
\\toprule
\\textbf{{Rank}} & \\textbf{{Gene Symbol}} & \\textbf{{ANOVA F-Score}} & \\textbf{{CV Selection Frequency}} \\\\
\\midrule
{genes_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width=0.68\\linewidth]{{results/shap_summary_random_forest.png}}
\\caption{{Random forest SHAP summary showing feature impact on proliferation class prediction.}}
\\end{{figure}}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width=0.68\\linewidth]{{results/pathway_enrichment.png}}
\\caption{{Enriched biological pathways (FDR < 0.05) representing transcriptomic cascade signatures downstream of cancer cell division.}}
\\end{{figure}}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width=0.68\\linewidth]{{results/clinical_dca.png}}
\\caption{{Clinical Decision Curve Analysis comparing Net Benefit of model-guided stratification vs. default intervention strategies.}}
\\end{{figure}}

\\begin{{table}}[H]
\\centering
\\caption{{Diagnostic performance and Number Needed to Treat (NNT) at clinical risk thresholds.}}
\\begin{{tabular}}{{lllllll}}
\\toprule
\\textbf{{Model}} & \\textbf{{Threshold}} & \\textbf{{Sensitivity}} & \\textbf{{Specificity}} & \\textbf{{PPV}} & \\textbf{{NPV}} & \\textbf{{NNT}} \\\\
\\midrule
{nnt_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

\\section{{Demographic Subgroups \\& Prognostic Validation}}
Subgroup analyses checked for performance differences across age, sex, and stage. Accuracy and ROC-AUC were consistent across groups (Table 6).

\\begin{{table}}[H]
\\centering
\\caption{{Subgroup demographic performance validation with interaction testing (Best model: Logistic Regression).}}
\\begin{{tabular}}{{lllllll}}
\\toprule
\\textbf{{Subgroup}} & \\textbf{{N}} & \\textbf{{Accuracy}} & \\textbf{{ROC-AUC}} & \\textbf{{ROC-AUC 95\\% CI}} & \\textbf{{Interaction p-val}} & \\textbf{{Interaction 95\\% CI}} \\\\
\\midrule
{sub_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

Kaplan-Meier curves compare survival by predicted proliferation class.

\\begin{{figure}}[H]
\\centering
\\includegraphics[width=0.62\\linewidth]{{results/kaplan_meier_geo.png}}
\\caption{{Kaplan-Meier overall survival curves comparing predicted high vs. low proliferation cohorts in GEO.}}
\\end{{figure}}

\\begin{{figure}}[H]
\\centering
\\includegraphics[width=0.68\\linewidth]{{results/kaplan_meier_stage_stratified.png}}
\\caption{{Kaplan-Meier overall survival curves stratified by stage and predicted proliferation class.}}
\\end{{figure}}

{cox_paragraph}

\\begin{{table}}[H]
\\centering
\\caption{{Multivariate Cox Proportional Hazards survival analysis summary.}}
\\begin{{tabular}}{{lllll}}
\\toprule
\\textbf{{Predictor / Covariate}} & \\textbf{{Coefficient}} & \\textbf{{Hazard Ratio (HR)}} & \\textbf{{95\\% CI}} & \\textbf{{p-value}} \\\\
\\midrule
{cox_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

\\section{{Pre-Processing Sensitivity \\& Robustness Analysis}}
Sensitivity analyses varied feature selection count (k) and variance threshold (VT). ROC-AUC remained stable across the tested ranges (Table 8).

\\begin{{table}}[H]
\\centering
\\caption{{Model pre-processing sensitivity analyses for feature selection size (k) and variance threshold (VT).}}
\\begin{{tabular}}{{lllll}}
\\toprule
\\textbf{{SelectKBest k}} & \\textbf{{Holdout ROC-AUC (k)}} & \\textbf{{Variance Threshold (VT)}} & \\textbf{{Features Passed VT}} & \\textbf{{Holdout ROC-AUC (VT)}} \\\\
\\midrule
{sens_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

\\section{{Discussion}}
{discussion}

\\begin{{table}}[H]
\\centering
\\caption{{Performance and methodological comparisons with published signatures.}}
\\begin{{tabular}}{{lllllll}}
\\toprule
\\textbf{{Study}} & \\textbf{{Year}} & \\textbf{{Cohort/Platform}} & \\textbf{{N}} & \\textbf{{AUC/Accuracy}} & \\textbf{{Leakage-controlled?}} & \\textbf{{Cross-platform validated?}} \\\\
\\midrule
{bench_content}
\\bottomrule
\\end{{tabular}}
\\end{{table}}

{discussion_pathway}

\\section{{Limitations and Future Directions}}
\\subsection*{{Limitations}}
\\begin{{itemize}}
    \\item Sample size: GEO GSE39582 (n=585) is moderate. CPTAC-COAD (n=105) has only 7 survival events.
    \\item Binarizing proliferation scores at the median is standard but arbitrary.
    \\item Microarray and RNA-seq have different dynamic ranges, requiring post-hoc calibration.
    \\item SHAP scores reflect correlation, not causation.
\\end{{itemize}}

\\subsection*{{Future Work}}
\\begin{{itemize}}
    \\item Prospective validation of the Top-3 Ensemble on new COAD biopsy cohorts.
    \\item qPCR knockdown of top SHAP genes (RPS3, RPS11) to test their role in growth regulation.
    \\item Integration of CNVs and somatic mutation data into the feature space.
\\end{{itemize}}

\\section{{References}}
\\begin{{enumerate}}
    \\item Marisa, L. et al. Gene expression classification of colon cancer into molecular subtypes: characterization, validation, and prognostic value. \\textit{{PLoS Medicine}}, 2013. DOI: 10.1371/journal.pmed.1001453
    \\item Whitfield, M. L. et al. Identification of genes periodically expressed in the human cell cycle and their expression in tumors. \\textit{{Molecular Biology of the Cell}}, 2002. DOI: 10.1091/mbc.02-02-0030
    \\item Lundberg, S. M. and Lee, S.-I. A unified approach to interpreting model predictions. \\textit{{Advances in Neural Information Processing Systems}}, 2017.
    \\item The Cancer Genome Atlas Research Network. TCGA Colon Adenocarcinoma data resource, accessed through the NCI Genomic Data Commons.
    \\item National Center for Biotechnology Information Gene Expression Omnibus. GSE39582 dataset record.
    \\item Zeng, D.-T. et al. Prognostic role of Ki-67 in colorectal carcinoma: Development and evaluation of machine learning prediction models. \\textit{{World Journal of Clinical Oncology}}, 2025. DOI: 10.5306/wjco.v16.i8.107306
    \\item Agesen, T. H. et al. ColoGuideEx: a robust gene classifier specific for stage II colorectal cancer prognosis. \\textit{{Gut}}, 2012. DOI: 10.1136/gutjnl-2011-301179
    \\item O'Connell, M. J. et al. Relationship between tumor gene expression and recurrence in four independent studies of patients with stage II/III colon cancer treated with surgery alone or surgery plus adjuvant fluorouracil plus leucovorin. \\textit{{Journal of Clinical Oncology}}, 2010. DOI: 10.1200/JCO.2010.28.9538
    \\item Langston, L. D. et al. Mcm10 promotes rapid isomerization of CMG-DNA for replisome bypass of lagging strand DNA blocks. \\textit{{eLife}}, 2017. DOI: 10.7554/eLife.29118
    \\item Bharadwaj, R., Qi, W., and Yu, H. Identification of two novel components of the human NDC80 kinetochore complex. \\textit{{Journal of Biological Chemistry}}, 2004. DOI: 10.1074/jbc.M310224200
    \\item Seipold, S. et al. Non-SMC condensin I complex proteins control chromosome segregation and survival of proliferating cells in the zebrafish neural retina. \\textit{{BMC Developmental Biology}}, 2009. DOI: 10.1186/1471-213X-9-40
    \\item Overmeer, R. M. et al. Replication factor C recruits DNA polymerase delta to sites of nucleotide excision repair but is not required for PCNA recruitment. \\textit{{Molecular and Cellular Biology}}, 2010. DOI: 10.1128/MCB.00285-10
\\end{{enumerate}}

\\section*{{Data and Code Availability}}
GEO GSE39582 is available at: \\url{{https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE39582}}. TCGA-COAD is available at UCSC Xena: \\url{{https://xenabrowser.net/}}. The repository code, trained pipelines, and reproducibility instructions are available at: \\url{{https://github.com/Ronisnotasianfr/ColoGrowth-ML}}.

\\section*{{Ethical Considerations and AI Disclosure}}
Secondary analysis of de-identified public datasets did not require institutional review board (IRB) approval. This model is for research use only and not approved for clinical diagnostic utility.

\\subsection*{{AI Assistance}}
Claude (Anthropic) was used as a coding assistant during implementation. All scientific decisions, study design, data interpretation, and conclusions are the author's own. Full prompt logs are available in the project repository.

\\end{{document}}
"""
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(tex_content)
    print(f"LaTeX paper written to {out_path}")


def main():
    parser = argparse.ArgumentParser(description="Build Word research paper from pipeline metrics")
    parser.add_argument("--dataset", default="geo", choices=["geo", "geo_pan", "tcga", "tcga_pan", "synthetic"])
    args = parser.parse_args()

    base = project_root()
    results_dir = base / "results"
    data_dir = base / "data" / "processed"
    out_docx = base / "paper" / "colon_cancer_growth_prediction_research_paper.docx"
    out_tex = base / "paper" / "colon_cancer_growth_prediction_research_paper.tex"

    metrics = load_leakage_fixed_metrics(args.dataset, results_dir)
    stats = load_dataset_stats(args.dataset, data_dir)

    doc = Document()
    apply_styles(doc)
    add_title_page(doc, metrics, stats)
    add_page_break(doc)
    add_methods(doc, stats)
    add_page_break(doc)
    add_results(doc, metrics, results_dir)
    add_page_break(doc)
    add_interpretation(doc, results_dir)
    add_page_break(doc)
    add_clinical_validation(doc, results_dir)
    add_page_break(doc)
    add_sensitivity_section(doc)
    add_page_break(doc)
    add_discussion(doc)
    add_page_break(doc)
    add_references(doc)
    doc.save(out_docx)
    print(f"Word paper written to {out_docx}")

    build_latex(metrics, stats, out_tex)


if __name__ == "__main__":
    main()
